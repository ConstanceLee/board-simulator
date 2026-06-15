import io
import os
import re
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from copy import deepcopy

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "woolworths_template.docx")
DOC_TITLE = "Board Simulation Report"
_HEADER_FILL = "EAF1FE"
_BORDER = "D9D9D9"
_BODY_AFTER = Pt(6)
_BULLET_AFTER = Pt(3)

def _fix_margins(document) -> None:
    for sect_pr in document.element.body.iter(qn("w:sectPr")):
        pg_mar = sect_pr.find(qn("w:pgMar"))
        if pg_mar is None:
            continue
        for attr in list(pg_mar.attrib):
            try:
                pg_mar.set(attr, str(int(round(float(pg_mar.get(attr))))))
            except (TypeError, ValueError):
                pass

def _clear_body(document) -> None:
    body = document.element.body
    for child in list(body):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(child)

def _fix_header_footer(document, date_str: str) -> None:
    sec = document.sections[0]
    replacements = {
        "My document title": DOC_TITLE,
        "Month 20XX": date_str,
    }
    for name in ("header", "first_page_header", "even_page_header"):
        for run in getattr(sec, name)._element.iter(qn("w:t")):
            if run.text in replacements:
                run.text = replacements[run.text]

    pg_num = sec._sectPr.find(qn("w:pgNumType"))
    if pg_num is not None:
        pg_num.set(qn("w:start"), "1")

    first_footer = sec.first_page_footer
    for para in list(first_footer.paragraphs):
        para._p.getparent().remove(para._p)
    first_footer._element.append(deepcopy(sec.footer.paragraphs[0]._p))

def _heading(document, text: str, level: int):
    p = document.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(20) if level == 2 else Pt(12)
    p.paragraph_format.space_after = _BODY_AFTER
    return p

def _label(document, label: str, value: str = "", *, italic: bool = False):
    p = document.add_paragraph(style="normal")
    p.paragraph_format.space_after = _BODY_AFTER
    run = p.add_run(label)
    run.bold = True
    run.italic = italic
    if value:
        p.add_run(value)
    return p

def _bullet(document, text: str, *, level: int = 0):
    p = document.add_paragraph(style="normal")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = _BULLET_AFTER
    p.add_run("◦  " if level else "•  ")
    
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)
    return p

def _set_cell(cell, text: str, *, bold: bool = False, fill: str | None = None) -> None:
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(
            tc_pr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): fill})
        )

def _table(document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(
            borders.makeelement(
                qn("w:" + edge),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "0",
                    qn("w:color"): _BORDER,
                },
            )
        )
    tbl_pr.append(borders)
    
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(tr_pr.makeelement(qn("w:cantSplit"), {}))
        
    for c, head in enumerate(headers):
        _set_cell(table.rows[0].cells[c], head, bold=True, fill=_HEADER_FILL)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            _set_cell(table.rows[r].cells[c], value)
    return table

def _parse_markdown_into_doc(document, md_text: str):
    """Parses a markdown block line-by-line and appends corresponding formatted elements."""
    lines = md_text.split('\n')
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
            
        if line_strip.startswith("###"):
            _heading(document, line_strip.replace("###", "").strip(), 3)
        elif line_strip.startswith("##"):
            _heading(document, line_strip.replace("##", "").strip(), 2)
        elif line_strip.startswith("#"):
            _heading(document, line_strip.replace("#", "").strip(), 1)
        elif line_strip.startswith("- ") or line_strip.startswith("* "):
            _bullet(document, line_strip[2:], level=0)
        elif line_strip.startswith("  - ") or line_strip.startswith("  * ") or line_strip.startswith("\t- ") or line_strip.startswith("\t* "):
            _bullet(document, line_strip.strip()[2:], level=1)
        elif re.match(r'^\d+\.\s', line_strip):
            _bullet(document, re.sub(r'^\d+\.\s', '', line_strip), level=0)
        else:
            p = document.add_paragraph(style="normal")
            p.paragraph_format.space_after = _BODY_AFTER
            parts = re.split(r'(\*\*[^*]+\*\*)', line_strip)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

def build_board_simulation_docx(board_paper_name: str, simulated_members: list, synthesis_markdown: str) -> bytes:
    doc = Document(TEMPLATE_PATH)
    _fix_margins(doc)
    
    now = datetime.now()
    date_str = f"{now.day} {now.strftime('%B %Y')}"
    _fix_header_footer(doc, date_str)
    _clear_body(doc)
    
    _heading(doc, "Woolworths Group Board Simulation Report", 1)
    _label(doc, "Simulation Date:", f" {date_str}")
    _label(doc, "Analysed Board Paper:", f" {board_paper_name}")
    
    # Executive Stances Table
    _heading(doc, "Executive Stances Summary", 2)
    roles_map = {
        "Scott Perkins": "Independent Chair of the Board",
        "Amanda Bardwell": "Chief Executive Officer (CEO)",
        "Maxine Brenner": "Non-Executive Director (Risk & Gov)",
        "Jennifer Carr-Smith": "Non-Executive Director (E-commerce)",
        "Philip Chronican": "Non-Executive Director (Governance)",
        "Kathee Tesija": "Non-Executive Director (Merchandising & Supply Chain)",
        "Warwick Bray": "Non-Executive Director (Audit & Finance Chair)",
        "Ken Meyer": "Non-Executive Director (Operations)",
        "Jon Alferness": "Non-Executive Director (Retail AI)"
    }
    
    headers = ["Board Member", "Role", "Simulated Stance"]
    rows = [[member.name, roles_map.get(member.name, "Non-Executive Director"), member.stance] for member in simulated_members]
    _table(doc, headers, rows)
    
    # Detailed responses
    _heading(doc, "Phase 2: Detailed Individual Responses", 2)
    for member in simulated_members:
        _heading(doc, member.name, 3)
        _label(doc, "Role:", f" {roles_map.get(member.name, 'Non-Executive Director')}", italic=True)
        _label(doc, "Stance:", f" {member.stance}")
        _label(doc, "Rationale:", f" {member.rationale}")
        _label(doc, "Questions & Concerns Raised:")
        for q in member.focus_points:
            _bullet(doc, q, level=0)
        _label(doc, "Key Data Request:", f" {member.key_request}")
        
    _parse_markdown_into_doc(doc, synthesis_markdown)
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
